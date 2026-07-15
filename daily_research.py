"""
Aria每日调研脚本 v2
覆盖：AI伴侣行业、技术论文、竞品监控、心理学研究
自动过滤无关内容，生成带启发总结的报告
"""
import time
import datetime
from pathlib import Path
from urllib.parse import quote

import requests
from bs4 import BeautifulSoup

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

REPORT_DIR = Path(r"D:\Aria\daily_reports")
REPORT_DIR.mkdir(exist_ok=True)

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
}

# ============================================
# 搜索配置：精确关键词 + 必须包含的词 + 排除的词
# ============================================
SEARCH_TOPICS = {
    "AI伴侣行业动态": {
        "queries": [
            '"AI伴侣" OR "AI陪伴" 融资 发布 2026',
            '"AI companion" app market funding startup 2026',
            'Character.AI OR Replika OR 星野 OR 猫箱 新功能 更新',
            'AI情感陪伴 用户增长 月活 市场规模',
        ],
        "must_contain": ["AI", "伴侣", "陪伴", "companion", "chatbot", "对话", "聊天"],
        "exclude": ["宠物", "猫粮", "狗粮", "抖音", "快手", "电商", "购物", "游戏攻略"],
        "trusted_domains": ["36kr.com", "zhihu.com", "techcrunch.com", "theverge.com", "aicompanionpick.com", "fortunebusinessinsights.com", "sohu.com", "sina.com", "163.com", "baidu.com"],
    },
    "技术论文追踪": {
        "queries": [
            'arxiv "emotional support" dialogue system 2026',
            'arxiv "persona consistency" LLM memory alignment',
            'arxiv "empathetic response" generation conversation',
            'arxiv "affective computing" emotion recognition dialogue',
        ],
        "must_contain": ["arxiv", "paper", "论文", "research", "LLM", "dialogue", "memory", "persona"],
        "exclude": ["宠物", "游戏", "电商", "购物"],
        "trusted_domains": ["arxiv.org", "ieee.org", "aclanthology.org", "semanticscholar.org", "openreview.net", "zhihu.com", "csdn.net", "github.com"],
    },
    "竞品监控": {
        "queries": [
            'Character.AI 2026 新功能 更新 changelog',
            'Replika AI 2026 update new features memory',
            '星野AI 2026 更新 功能 月活',
            '猫箱AI 2026 更新 功能 字节跳动',
        ],
        "must_contain": ["Character.AI", "Replika", "星野", "猫箱", "AI", "伴侣", "陪伴"],
        "exclude": ["宠物", "电商", "购物", "游戏攻略"],
        "trusted_domains": ["36kr.com", "zhihu.com", "aicompanionpick.com", "sohu.com", "sina.com", "163.com", "crsky.com"],
    },
    "心理学与人格研究": {
        "queries": [
            'MBTI 恋爱 配对 适配 心理学研究 2026',
            '依恋类型 亲密关系 心理学 attachment style',
            '九型人格 恋爱 关系 互动模式',
            'human-AI relationship trust psychology companion',
        ],
        "must_contain": ["MBTI", "人格", "心理", "依恋", "恋爱", "配对", "personality", "attachment"],
        "exclude": ["宠物", "电商", "购物", "游戏"],
        "trusted_domains": ["zhihu.com", "16personalities.com", "psyctest.cn", "webseo9.com", "psychologytoday.com", "apa.org"],
    },
}

# 启发规则：关键词 -> 启发内容
INSPIRATION_RULES = [
    (["memory", "记忆", "长期记忆", "long-term"], "记忆系统：发现新记忆架构思路"),
    (["persona", "人格", "人设", "consistency"], "人格一致性：关注防崩塌技术"),
    (["emotion", "情感", "情绪", "empathy", "共情"], "情感计算：新情感识别/共情技术"),
    (["social", "社交", "群聊", "多角色", "multi-character"], "社交功能：AI间互动新进展"),
    (["memory dashboard", "记忆可视化"], "记忆可视化：增强用户信任"),
    (["regulation", "监管", "合规", "未成年人"], "合规提醒：新监管政策"),
    (["compatibility", "匹配", "配对", "依恋"], "适配算法：新匹配理论"),
    (["灵魂", "soul", "壁垒", "context"], "核心壁垒：记忆是竞争力"),
    (["subscription", "付费", "monetization"], "商业化：新付费模式参考"),
]


def search_bing(query, max_results=5):
    """Bing搜索，返回结果列表"""
    try:
        url = f"https://www.bing.com/search?q={quote(query)}&count={max_results}"
        resp = requests.get(url, headers=HEADERS, timeout=15)
        soup = BeautifulSoup(resp.text, "html.parser")
        results = []
        for item in soup.select("li.b_algo")[:max_results]:
            title_el = item.select_one("h2 a")
            snippet_el = item.select_one(".b_caption p")
            if title_el:
                results.append({
                    "title": title_el.get_text(strip=True),
                    "snippet": snippet_el.get_text(strip=True)[:300] if snippet_el else "",
                    "url": title_el.get("href", ""),
                })
        return results
    except Exception as e:
        return [{"title": f"搜索失败: {e}", "snippet": "", "url": ""}]


def is_relevant(result, must_contain, exclude, trusted_domains=None):
    """判断搜索结果是否相关"""
    text = (result.get("title", "") + " " + result.get("snippet", "")).lower()
    url = result.get("url", "").lower()

    # 排除无关内容
    for word in exclude:
        if word.lower() in text:
            return False

    # 排除明显无关的网站（词典、百科、工具站）
    junk_domains = ["iciba.com", "youdao.com", "zdic.net", "dict.", "cambridge.org", "microsoft.com/translator", "ai-bot.cn", "ai-kit.cn", "runoob.com"]
    for domain in junk_domains:
        if domain in url:
            return False

    # 如果有白名单，优先匹配白名单
    if trusted_domains:
        for domain in trusted_domains:
            if domain in url:
                return True
        # 不在白名单的结果，仍需满足关键词匹配
        for word in must_contain:
            if word.lower() in text:
                return True
        return False

    # 无白名单时，按关键词匹配
    for word in must_contain:
        if word.lower() in text:
            return True
    return False


def research_topic(topic_name, config):
    """调研一个主题"""
    print(f"  搜索: {topic_name}...")
    all_results = []
    for query in config["queries"]:
        print(f"    关键词: {query[:40]}...")
        results = search_bing(query, max_results=5)
        for r in results:
            if is_relevant(r, config["must_contain"], config["exclude"], config.get("trusted_domains")):
                all_results.append(r)
        time.sleep(2)
    return all_results


def extract_inspiration(all_data):
    """从搜索结果提取启发"""
    all_text = ""
    for results in all_data.values():
        for r in results:
            all_text += r.get("title", "") + " " + r.get("snippet", "") + " "

    inspirations = []
    for keywords, insight in INSPIRATION_RULES:
        for kw in keywords:
            if kw.lower() in all_text.lower():
                inspirations.append(insight)
                break

    # 去重保持顺序
    seen = set()
    unique = []
    for item in inspirations:
        if item not in seen:
            seen.add(item)
            unique.append(item)

    return unique or ["今日暂无明显新启发，继续积累。"]


def deduplicate(results):
    """按标题去重"""
    seen = set()
    unique = []
    for r in results:
        key = r["title"][:50]
        if key and key not in seen:
            seen.add(key)
            unique.append(r)
    return unique


def generate_markdown(date_str, all_data):
    """生成Markdown报告"""
    lines = [f"# Aria每日调研报告", f"日期：{date_str}", ""]

    total = 0
    for topic, results in all_data.items():
        lines.append(f"## {topic}")
        unique = deduplicate(results)
        lines.append(f"共找到 {len(unique)} 条相关信息")
        lines.append("")
        for i, r in enumerate(unique[:6], 1):
            lines.append(f"### {i}. {r['title']}")
            if r["snippet"]:
                lines.append(r["snippet"])
            if r["url"]:
                lines.append(f"链接: {r['url']}")
            lines.append("")
            total += 1
        lines.append("---")
        lines.append("")

    # 启发总结
    lines.append("## 对Aria的启发与新思路")
    lines.append("")
    for item in extract_inspiration(all_data):
        lines.append(f"- {item}")
    lines.append("")

    lines.append(f"## 总结")
    lines.append(f"今日共搜索 {sum(len(v) for v in all_data.values())} 条，有效 {total} 条。")
    lines.append(f"报告生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    return "\n".join(lines)


def save_docx(date_str, all_data):
    """生成Word报告"""
    if not HAS_DOCX:
        return
    doc = Document()
    doc.add_heading('Aria每日调研报告', level=0)
    doc.add_paragraph(f'日期：{date_str}')

    for topic, results in all_data.items():
        doc.add_heading(topic, level=1)
        for r in deduplicate(results)[:6]:
            doc.add_heading(r["title"], level=2)
            if r["snippet"]:
                doc.add_paragraph(r["snippet"])
            if r["url"]:
                doc.add_paragraph(f'链接: {r["url"]}')

    doc.add_heading('对Aria的启发与新思路', level=1)
    for item in extract_inspiration(all_data):
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(f'报告生成时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    path = REPORT_DIR / f"research_{date_str}.docx"
    doc.save(str(path))
    print(f"  Word报告: {path}")


def main():
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    print(f"{'='*50}")
    print(f"Aria每日调研 - {date_str}")
    print(f"{'='*50}")

    all_data = {}
    for topic, config in SEARCH_TOPICS.items():
        results = research_topic(topic, config)
        all_data[topic] = results
        print(f"  {topic}: {len(results)} 条有效")

    md = generate_markdown(date_str, all_data)
    md_path = REPORT_DIR / f"research_{date_str}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    print(f"\nMarkdown报告: {md_path}")

    save_docx(date_str, all_data)
    print(f"\n完成！报告目录: {REPORT_DIR}")


if __name__ == "__main__":
    main()
